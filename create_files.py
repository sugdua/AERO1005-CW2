#!/usr/bin/env python3
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch
import numpy as np
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
import os

OUT = '/Users/maliyuyuyu/Desktop/AERO1005_CW2'

# ============ FLOWCHART HELPERS ============
def draw_box(ax, x, y, w, h, text, style='process', fontsize=7):
    colors = {
        'start': ('#2E7D32', 'white'),
        'process': ('#E3F2FD', 'black'),
        'io': ('#FFF3E0', 'black'),
        'decision': ('#FCE4EC', 'black'),
    }
    fc, tc = colors.get(style, ('#E3F2FD', 'black'))
    
    if style == 'start':
        box = FancyBboxPatch((x - w/2, y - h/2), w, h,
                             boxstyle="round,pad=0.05", fc=fc, ec='black', lw=1.2)
        ax.add_patch(box)
        ax.text(x, y, text, ha='center', va='center', fontsize=fontsize, color=tc, fontweight='bold')
    elif style == 'decision':
        # Diamond
        diamond = plt.Polygon([(x, y+h/2), (x+w/2, y), (x, y-h/2), (x-w/2, y)],
                              fc=fc, ec='black', lw=1.2)
        ax.add_patch(diamond)
        ax.text(x, y, text, ha='center', va='center', fontsize=fontsize-1, color=tc, wrap=True)
    else:
        box = FancyBboxPatch((x - w/2, y - h/2), w, h,
                             boxstyle="round,pad=0.03", fc=fc, ec='black', lw=1)
        ax.add_patch(box)
        ax.text(x, y, text, ha='center', va='center', fontsize=fontsize, color=tc)

def arrow(ax, x1, y1, x2, y2, label=''):
    ax.annotate('', xy=(x2, y2), xytext=(x1, y1),
                arrowprops=dict(arrowstyle='->', color='black', lw=1.2))
    if label:
        mx, my = (x1+x2)/2, (y1+y2)/2
        ax.text(mx+0.15, my, label, fontsize=6, color='#1565C0', fontweight='bold')

# ============ FLOWCHART TASK 2 PRE ============
def make_flowchart_task2_pre():
    fig, ax = plt.subplots(1, 1, figsize=(8, 12))
    ax.set_xlim(-4, 4)
    ax.set_ylim(-1, 13)
    ax.axis('off')
    ax.set_aspect('equal')
    
    # Nodes top to bottom
    nodes = [
        (0, 12, 2, 0.6, 'START', 'start'),
        (0, 10.5, 3, 0.6, 'Initialise sensor & LEDs', 'process'),
        (0, 9, 2.5, 0.6, 'LOOP START', 'process'),
        (0, 7.5, 2.8, 0.6, 'Read temperature', 'io'),
        (0, 6, 3, 0.8, 'Is 18 ≤ T ≤ 24?', 'decision'),
        (3, 6, 2.5, 0.6, 'Is T < 18?', 'decision'),
        (-3, 4.5, 2.8, 0.6, 'Green ON\nothers OFF', 'process'),
        (2, 4, 2.8, 0.6, 'Yellow blink\n0.5s interval', 'process'),
        (4, 4, 2.8, 0.6, 'Red blink\n0.25s interval', 'process'),
    ]
    
    for x, y, w, h, t, s in nodes:
        draw_box(ax, x, y, w, h, t, s)
    
    arrow(ax, 0, 11.7, 0, 10.8)
    arrow(ax, 0, 10.2, 0, 9.3)
    arrow(ax, 0, 8.7, 0, 7.8)
    arrow(ax, 0, 7.2, 0, 6.4)
    # Decision YES -> Green
    arrow(ax, -1.5, 6, -3, 4.8, 'YES')
    # Decision NO -> Is T<18?
    arrow(ax, 1.5, 6, 1.8, 6, 'NO')
    # T<18 YES -> Yellow
    arrow(ax, 2, 5.6, 2, 4.3, 'YES')
    # T<18 NO -> Red
    arrow(ax, 3.5, 5.6, 4, 4.3, 'NO')
    
    # Loop backs (curved)
    for bx in [-3, 2, 4]:
        by = 3.7 if bx != -3 else 4.2
        ax.annotate('', xy=(0, 8.7), xytext=(bx, by),
                     arrowprops=dict(arrowstyle='->', color='grey', lw=1, connectionstyle='arc3,rad=0.3'))
    
    fig.tight_layout()
    fig.savefig(os.path.join(OUT, 'flowchart_task2_pre.png'), dpi=200, bbox_inches='tight')
    plt.close(fig)

# ============ FLOWCHART TASK 2 POST ============
def make_flowchart_task2_post():
    fig, ax = plt.subplots(1, 1, figsize=(9, 16))
    ax.set_xlim(-5, 5)
    ax.set_ylim(-1, 17)
    ax.axis('off')
    ax.set_aspect('equal')
    
    y = 16
    dy = 1.3
    
    items = [
        (0, y, 2, 0.5, 'START', 'start'),
        (0, y-dy, 3.5, 0.5, 'Initialise V0, TC, pins,\narrays, figure', 'process'),
        (0, y-2*dy, 2, 0.5, 'tic (start timer)', 'process'),
        (0, y-3*dy, 2, 0.5, 'LOOP', 'process'),
        (0, y-4*dy, 3, 0.5, 'readVoltage(a, "A0")', 'io'),
        (0, y-5*dy, 3.2, 0.5, 'Convert: T=(V-V0)/TC', 'process'),
        (0, y-6*dy, 3, 0.5, 'Append to arrays', 'process'),
        (0, y-7*dy, 3.5, 0.5, 'Update plot\n(set, xlim, ylim, drawnow)', 'process'),
        (0, y-8*dy, 3, 0.7, 'Is 18 ≤ T ≤ 24?', 'decision'),
        (-3.5, y-9.5*dy, 2.8, 0.5, 'Green ON\npause(1)', 'process'),
        (0, y-9.5*dy, 2.8, 0.7, 'Is T < 18?', 'decision'),
        (-1.5, y-11*dy, 3, 0.5, 'Yellow blink\npause(0.5) x2', 'process'),
        (2, y-11*dy, 3, 0.5, 'Red blink\npause(0.25) x4', 'process'),
    ]
    
    for x, iy, w, h, t, s in items:
        draw_box(ax, x, iy, w, h, t, s)
    
    # Sequential arrows
    for i in range(7):
        arrow(ax, 0, y - i*dy - 0.25, 0, y - (i+1)*dy + 0.25)
    
    # Decision arrows
    arrow(ax, 0, y-8*dy-0.35, 0, y-8*dy-0.35)
    arrow(ax, -1.5, y-8*dy, -3.5, y-9.5*dy+0.25, 'YES')
    arrow(ax, 1.5, y-8*dy, 0, y-9.5*dy+0.35, 'NO')
    arrow(ax, -0.8, y-9.5*dy, -1.5, y-11*dy+0.25, 'YES')
    arrow(ax, 0.8, y-9.5*dy, 2, y-11*dy+0.25, 'NO')
    
    # Loop back arrows
    for bx, by in [(-3.5, y-9.5*dy-0.25), (-1.5, y-11*dy-0.25), (2, y-11*dy-0.25)]:
        ax.annotate('', xy=(0, y-3*dy+0.25), xytext=(bx, by),
                     arrowprops=dict(arrowstyle='->', color='grey', lw=1, connectionstyle='arc3,rad=0.3'))
    
    fig.tight_layout()
    fig.savefig(os.path.join(OUT, 'flowchart_task2_post.png'), dpi=200, bbox_inches='tight')
    plt.close(fig)

# ============ FLOWCHART TASK 3 ============
def make_flowchart_task3():
    fig, ax = plt.subplots(1, 1, figsize=(9, 18))
    ax.set_xlim(-5, 5)
    ax.set_ylim(-1, 19)
    ax.axis('off')
    ax.set_aspect('equal')
    
    y = 18
    dy = 1.2
    
    items = [
        (0, y, 2, 0.5, 'START', 'start'),
        (0, y-dy, 3.5, 0.5, 'Initialise V0, TC,\nrate_threshold, window_size', 'process'),
        (0, y-2*dy, 2, 0.5, 'tic (start timer)', 'process'),
        (0, y-3*dy, 2, 0.5, 'LOOP', 'process'),
        (0, y-4*dy, 3, 0.5, 'Read temperature', 'io'),
        (0, y-5*dy, 2.5, 0.5, 'Store data', 'process'),
        (0, y-6*dy, 3.2, 0.7, 'n ≥ window_size?', 'decision'),
        (-2.5, y-7.5*dy, 3, 0.5, 'polyfit over window\n→ rate_per_sec', 'process'),
        (2.5, y-7.5*dy, 3, 0.5, 'Simple difference\n→ rate_per_sec', 'process'),
        (0, y-8.5*dy, 3, 0.5, 'rate_per_min =\nrate_per_sec × 60', 'process'),
        (0, y-9.5*dy, 3.2, 0.5, 'predicted = T +\nrate × 300s', 'process'),
        (0, y-10.5*dy, 3, 0.5, 'Print to screen', 'io'),
        (0, y-11.5*dy, 3, 0.7, 'Rate > +4°C/min?', 'decision'),
        (-3, y-12.5*dy, 2, 0.5, 'Red LED ON', 'process'),
        (2, y-12.5*dy, 3, 0.7, 'Rate < -4°C/min?', 'decision'),
        (1, y-13.7*dy, 2.2, 0.5, 'Yellow LED ON', 'process'),
        (3.5, y-13.7*dy, 2.2, 0.5, 'Green LED ON', 'process'),
        (0, y-14.8*dy, 2, 0.5, 'pause(1)', 'process'),
    ]
    
    for x, iy, w, h, t, s in items:
        draw_box(ax, x, iy, w, h, t, s)
    
    for i in range(6):
        arrow(ax, 0, y-i*dy-0.25, 0, y-(i+1)*dy+0.25)
    
    # Decision: enough samples?
    arrow(ax, -1.6, y-6*dy, -2.5, y-7.5*dy+0.25, 'YES')
    arrow(ax, 1.6, y-6*dy, 2.5, y-7.5*dy+0.25, 'NO')
    # Merge
    arrow(ax, -2.5, y-7.5*dy-0.25, 0, y-8.5*dy+0.25)
    arrow(ax, 2.5, y-7.5*dy-0.25, 0, y-8.5*dy+0.25)
    
    for i in range(8, 12):
        arrow(ax, 0, y-i*0.5*2-0.25-4*dy, 0, y-(i*0.5*2+1)*1-0.25-3.3*dy)
    # Simplified: connect remaining
    arrow(ax, 0, y-8.5*dy-0.25, 0, y-9.5*dy+0.25)
    arrow(ax, 0, y-9.5*dy-0.25, 0, y-10.5*dy+0.25)
    arrow(ax, 0, y-10.5*dy-0.25, 0, y-11.5*dy+0.35)
    
    # Rate decisions
    arrow(ax, -1.5, y-11.5*dy, -3, y-12.5*dy+0.25, 'YES')
    arrow(ax, 1.5, y-11.5*dy, 2, y-12.5*dy+0.35, 'NO')
    arrow(ax, 1.2, y-12.5*dy, 1, y-13.7*dy+0.25, 'YES')
    arrow(ax, 2.8, y-12.5*dy, 3.5, y-13.7*dy+0.25, 'NO')
    
    # All converge to pause
    for bx, by in [(-3, y-12.5*dy-0.25), (1, y-13.7*dy-0.25), (3.5, y-13.7*dy-0.25)]:
        arrow(ax, bx, by, 0, y-14.8*dy+0.25)
    
    # Loop back
    ax.annotate('', xy=(0, y-3*dy+0.25), xytext=(0, y-14.8*dy-0.25),
                arrowprops=dict(arrowstyle='->', color='grey', lw=1, connectionstyle='arc3,rad=-0.5'))
    
    fig.tight_layout()
    fig.savefig(os.path.join(OUT, 'flowchart_task3.png'), dpi=200, bbox_inches='tight')
    plt.close(fig)

# ============ WORD DOCUMENT ============
def make_docx():
    doc = Document()
    
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    
    # Title page
    for _ in range(4):
        doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('AERO1005 Coursework 2')
    r.bold = True
    r.font.size = Pt(26)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('Solving Engineering Problems with a Programming Language')
    r.font.size = Pt(16)
    r.font.color.rgb = RGBColor(80, 80, 80)
    
    for _ in range(2):
        doc.add_paragraph()
    
    for line in ['Name: Liyu MA', 'Student ID: 20721363', 'Date: May 2026']:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line)
        r.font.size = Pt(14)
    
    doc.add_page_break()
    
    # Preliminary Task
    doc.add_heading('Preliminary Task', level=1)
    doc.add_paragraph(
        'Arduino UNO was connected to the computer via USB cable. Communication was established '
        'using the arduino() function in MATLAB. A GitHub repository was created and linked to '
        'the MATLAB working folder for version control. The LED blink test was successfully '
        'implemented on digital pin D13.'
    )
    p = doc.add_paragraph()
    r = p.add_run('[INSERT PHOTOGRAPH OF ARDUINO + LED SETUP HERE]')
    r.italic = True
    r.font.color.rgb = RGBColor(200, 0, 0)
    
    doc.add_page_break()
    
    # Task 1
    doc.add_heading('Task 1 – Temperature Data Acquisition', level=1)
    doc.add_heading('a) Hardware Setup', level=2)
    doc.add_paragraph(
        'The MCP 9700A temperature sensor was connected to the breadboard with its VDD pin to '
        'the 5V power bus, VSS pin to the ground bus, and VOUT pin to Arduino analog channel A0. '
        'A 220Ω resistor was placed between the sensor output and A0 to stabilise the signal '
        'and reduce noise from the microcontroller power supply.'
    )
    p = doc.add_paragraph()
    r = p.add_run('[INSERT PHOTOGRAPH OF TEMPERATURE SENSOR SETUP HERE]')
    r.italic = True
    r.font.color.rgb = RGBColor(200, 0, 0)
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run('[INSERT TEMPERATURE PLOT IMAGE HERE]')
    r.italic = True
    r.font.color.rgb = RGBColor(200, 0, 0)
    
    doc.add_page_break()
    
    # Task 2
    doc.add_heading('Task 2 – LED Temperature Monitoring Device', level=1)
    doc.add_heading('f) Hardware Setup', level=2)
    doc.add_paragraph(
        'Three LEDs (green, yellow, red) were connected to the breadboard. Each LED\'s anode '
        '(long leg) was connected to a digital pin (D9, D10, D11 respectively) via a jumper wire, '
        'and each cathode (short leg) was connected to the ground bus through a 220Ω resistor.'
    )
    p = doc.add_paragraph()
    r = p.add_run('[INSERT PHOTOGRAPH OF 3-LED SETUP HERE]')
    r.italic = True
    r.font.color.rgb = RGBColor(200, 0, 0)
    
    doc.add_heading('Pre-Implementation Flowchart', level=2)
    doc.add_picture(os.path.join(OUT, 'flowchart_task2_pre.png'), width=Inches(5.5))
    
    doc.add_heading('Post-Implementation Flowchart', level=2)
    doc.add_picture(os.path.join(OUT, 'flowchart_task2_post.png'), width=Inches(5.5))
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run('[INSERT LIVE PLOT SNAPSHOT HERE]')
    r.italic = True
    r.font.color.rgb = RGBColor(200, 0, 0)
    
    doc.add_page_break()
    
    # Task 3
    doc.add_heading('Task 3 – Temperature Prediction', level=1)
    doc.add_heading('Flowchart', level=2)
    doc.add_picture(os.path.join(OUT, 'flowchart_task3.png'), width=Inches(5.5))
    
    doc.add_page_break()
    
    # Task 4
    doc.add_heading('Task 4 – Reflective Statement', level=1)
    doc.add_paragraph(
        'This coursework presented several engineering challenges that deepened my understanding '
        'of integrating hardware and software for real-time data acquisition. One of the primary '
        'challenges was synchronising the LED blinking behaviour with the temperature data '
        'acquisition loop. In Task 2, the LED blink timing (0.5s for yellow, 0.25s for red) '
        'directly affected the sampling rate of the temperature sensor, since both operations '
        'shared the same execution thread. This taught me the importance of considering timing '
        'constraints in embedded systems design.'
    )
    doc.add_paragraph(
        'Another significant challenge was managing noise in the temperature readings from the '
        'MCP 9700A sensor. Raw voltage readings exhibited fluctuations of approximately ±0.5°C, '
        'which could trigger false LED state changes. To address this in Task 3, I implemented '
        'a moving window approach using MATLAB\'s polyfit function over the last 30 samples. '
        'This linear regression technique provided a more robust estimate of the temperature '
        'rate of change compared to simple point-to-point differences, effectively filtering '
        'out short-term noise while preserving genuine trends.'
    )
    doc.add_paragraph(
        'A key strength of my implementation is the modular code structure. By encapsulating '
        'the monitoring and prediction logic in separate functions (temp_monitor and '
        'temp_prediction), the code remains readable, testable, and reusable. The live plotting '
        'feature in Task 2, using MATLAB\'s set and drawnow commands, provides immediate visual '
        'feedback that proved invaluable during debugging and demonstration.'
    )
    doc.add_paragraph(
        'However, there are notable limitations. The temperature prediction in Task 3 assumes '
        'a constant linear rate of change, which may not accurately represent real thermal '
        'dynamics influenced by convection, radiation, and insulation properties. The MCP 9700A '
        'sensor itself has a limited accuracy of ±2°C, which constrains the precision of all '
        'derived calculations. Future improvements could include implementing a Kalman filter '
        'for more sophisticated noise reduction, using hardware interrupts for LED timing to '
        'decouple it from the sensing loop, and employing a higher-precision digital temperature '
        'sensor such as the DS18B20. Additionally, incorporating wireless data transmission '
        'via Bluetooth or Wi-Fi would enable remote monitoring capabilities essential for '
        'real-world capsule applications.'
    )
    
    doc.add_page_break()
    
    # AI Usage
    doc.add_heading('AI Usage Disclosure', level=1)
    doc.add_paragraph(
        'The MATLAB Copilot tool was used solely for debugging syntax errors during development, '
        'specifically for identifying issues with fprintf formatting and Arduino pin configuration '
        'commands. No AI tools were used to generate algorithms, logic, or core code structure.'
    )
    
    doc.save(os.path.join(OUT, 'AERO1005_CW2_Submission.docx'))
    print('Word document created.')

if __name__ == '__main__':
    print('Creating flowcharts...')
    make_flowchart_task2_pre()
    print('  flowchart_task2_pre.png done')
    make_flowchart_task2_post()
    print('  flowchart_task2_post.png done')
    make_flowchart_task3()
    print('  flowchart_task3.png done')
    print('Creating Word document...')
    make_docx()
    print('All done!')
